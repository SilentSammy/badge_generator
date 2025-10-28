from docx import Document
from docx.shared import Inches
import shutil
import os

def copy_and_replace_template(template_path, output_path, old_text, new_text):
    """
    Copy a Word document template and replace specified text.
    
    Args:
        template_path (str): Path to the template document
        output_path (str): Path for the new document
        old_text (str): Text to replace
        new_text (str): Replacement text
    """
    try:
        # Copy the template file
        shutil.copy2(template_path, output_path)
        print(f"Template copied to: {output_path}")
        
        # Open the copied document
        doc = Document(output_path)
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
        
        # Replace text in tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
        
        # Save the modified document
        doc.save(output_path)
        print(f"Text replacement completed. '{old_text}' replaced with '{new_text}'")
        
    except FileNotFoundError:
        print(f"Error: Template file '{template_path}' not found.")
    except Exception as e:
        print(f"Error: {str(e)}")

def replace_text_with_image(doc_path, output_path, text_to_replace, image_path, width_inches=2.0):
    """
    Replace specified text with an image in a Word document.
    
    Args:
        doc_path (str): Path to the input document
        output_path (str): Path to save the modified document
        text_to_replace (str): Text to replace with image
        image_path (str): Path to the replacement image
        width_inches (float): Width of the image in inches
    """
    try:
        if not os.path.exists(image_path):
            print(f"Error: Image file '{image_path}' not found.")
            return False
        
        doc = Document(doc_path)
        
        # Search through all paragraphs
        for paragraph in doc.paragraphs:
            if text_to_replace in paragraph.text:
                # Clear the paragraph
                paragraph.clear()
                # Add the image
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                print(f"Replaced '{text_to_replace}' with image: {image_path}")
        
        # Search through table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if text_to_replace in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(width_inches))
                            print(f"Replaced '{text_to_replace}' in table with image: {image_path}")
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error replacing text with image: {str(e)}")
        return False

def replace_existing_image(doc_path, output_path, new_image_path):
    """
    Option 2: Replace the existing embedded image with a new one, maintaining dimensions and position.
    
    Args:
        doc_path (str): Path to the input document
        output_path (str): Path to save the modified document  
        new_image_path (str): Path to the new image file
    """
    try:
        if not os.path.exists(new_image_path):
            print(f"Error: New image file '{new_image_path}' not found.")
            return False
        
        # Copy the document first
        if doc_path != output_path:
            shutil.copy2(doc_path, output_path)
        
        doc = Document(output_path)
        
        # Find the existing image and replace it
        image_replaced = False
        
        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                # Check for images in this run
                inline_shapes = run._element.xpath('.//a:blip')
                if inline_shapes and not image_replaced:
                    print(f"Found existing image in paragraph {p_idx}, run {r_idx}")
                    
                    # Get current dimensions
                    drawings = run._element.xpath('.//w:drawing')
                    current_width = None
                    current_height = None
                    
                    for drawing in drawings:
                        inlines = drawing.xpath('.//wp:inline')
                        for inline in inlines:
                            extent = inline.xpath('.//wp:extent')[0]
                            cx = int(extent.get('cx'))  # Width in EMUs
                            cy = int(extent.get('cy'))  # Height in EMUs
                            current_width = cx / 914400  # Convert to inches
                            current_height = cy / 914400
                            break
                    
                    # Remove the existing run with image
                    paragraph._element.remove(run._element)
                    
                    # Add new image with same dimensions
                    new_run = paragraph.add_run()
                    if current_width:
                        new_run.add_picture(new_image_path, width=Inches(current_width))
                        print(f"Replaced image with {new_image_path} (maintained {current_width:.2f}\" x {current_height:.2f}\")")
                    else:
                        new_run.add_picture(new_image_path, width=Inches(5.33))  # Fallback
                        print(f"Replaced image with {new_image_path} (fallback dimensions)")
                    
                    image_replaced = True
                    break
            
            if image_replaced:
                break
        
        if not image_replaced:
            print("No existing image found to replace.")
            return False
        
        # Save the document
        doc.save(output_path)
        print(f"Document saved as: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error replacing existing image: {str(e)}")
        return False

def replace_all_images_with_dummies(doc_path, output_path, dummy_dir="dummy_images"):
    """
    Replace ALL images in document with numbered dummy images in order.
    
    Args:
        doc_path (str): Path to input document
        output_path (str): Path to save modified document
        dummy_dir (str): Directory containing dummy images
    """
    try:
        if doc_path != output_path:
            shutil.copy2(doc_path, output_path)
        
        doc = Document(output_path)
        
        # Get list of dummy images
        dummy_files = sorted([f for f in os.listdir(dummy_dir) if f.endswith('.jpg')])
        if not dummy_files:
            print(f"No dummy images found in {dummy_dir}")
            return False
        
        print(f"Found {len(dummy_files)} dummy images")
        
        image_count = 0
        
        # Check paragraphs for images
        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                # Check for both modern and legacy image formats
                blips = run._element.xpath('.//a:blip')
                picts = run._element.xpath('.//w:pict')
                
                if blips or picts:
                    if image_count < len(dummy_files):
                        dummy_file = os.path.join(dummy_dir, dummy_files[image_count])
                        
                        # Get current dimensions if possible
                        current_width = None
                        try:
                            drawings = run._element.xpath('.//w:drawing')
                            for drawing in drawings:
                                inlines = drawing.xpath('.//wp:inline') + drawing.xpath('.//wp:anchor')
                                for inline in inlines:
                                    extents = inline.xpath('.//wp:extent')
                                    if extents:
                                        cx = int(extents[0].get('cx', 0))
                                        if cx > 0:
                                            current_width = cx / 914400
                                            break
                        except:
                            pass
                        
                        # Remove the existing run
                        paragraph._element.remove(run._element)
                        
                        # Add new dummy image
                        new_run = paragraph.add_run()
                        width = Inches(current_width) if current_width else Inches(2.0)
                        new_run.add_picture(dummy_file, width=width)
                        
                        image_count += 1
                        print(f"Image {image_count}: Replaced paragraph image with {dummy_files[image_count-1]}")
        
        # Check tables for images
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_num = row_idx * len(table.columns) + col_idx + 1
                    
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            # Check for both modern and legacy image formats
                            blips = run._element.xpath('.//a:blip')
                            picts = run._element.xpath('.//w:pict')
                            
                            if blips or picts:
                                if image_count < len(dummy_files):
                                    dummy_file = os.path.join(dummy_dir, dummy_files[image_count])
                                    
                                    # Get current dimensions if possible
                                    current_width = None
                                    try:
                                        drawings = run._element.xpath('.//w:drawing')
                                        for drawing in drawings:
                                            inlines = drawing.xpath('.//wp:inline') + drawing.xpath('.//wp:anchor')
                                            for inline in inlines:
                                                extents = inline.xpath('.//wp:extent')
                                                if extents:
                                                    cx = int(extents[0].get('cx', 0))
                                                    if cx > 0:
                                                        current_width = cx / 914400
                                                        break
                                    except:
                                        pass
                                    
                                    # Remove the existing run
                                    paragraph._element.remove(run._element)
                                    
                                    # Add new dummy image
                                    new_run = paragraph.add_run()
                                    width = Inches(current_width) if current_width else Inches(2.0)
                                    new_run.add_picture(dummy_file, width=width)
                                    
                                    image_count += 1
                                    print(f"Image {image_count}: Replaced cell {cell_num} image with {dummy_files[image_count-1]}")
        
        doc.save(output_path)
        print(f"\nSuccessfully replaced {image_count} images in {output_path}")
        return True
        
    except Exception as e:
        print(f"Error replacing images: {str(e)}")
        return False

if __name__ == "__main__":
    # Test image replacement with Test.docx
    test_file = "Test.docx"
    output_file = "Test_with_dummies.docx"
    
    print("=== Testing Image Detection and Replacement ===")
    print(f"Input file: {test_file}")
    print(f"Output file: {output_file}")
    
    # Replace all images with dummy images
    success = replace_all_images_with_dummies(test_file, output_file)
    
    if success:
        print(f"\n✅ Successfully created {output_file}")
        print("Check the document to see if the images were replaced with numbered dummies!")
    else:
        print(f"\n❌ Failed to process {test_file}")