from docx import Document
from docx.shared import Inches
import shutil
import os

def replace_text_placeholders(doc_path, output_path, text_replacements):
    """
    Replace text placeholders (like text1, text2, etc.) with multiline text.
    Handles \n characters by creating proper paragraph breaks.
    
    Args:
        doc_path (str): Path to input document
        output_path (str): Path to save modified document
        text_replacements (dict): Dictionary with placeholders as keys and replacement text as values
                                 Format: {'text1': 'John Smith\nSoftware Engineer\n123 Main St'}
    """
    try:
        # Copy the document
        if doc_path != output_path:
            shutil.copy2(doc_path, output_path)
        
        doc = Document(output_path)
        
        if not doc.tables:
            print("No tables found in document")
            return False
        
        table = doc.tables[0]  # Get first table
        total_cells = len(table.rows) * len(table.columns)
        
        print(f"Table has {len(table.rows)} rows x {len(table.columns)} columns = {total_cells} cells")
        
        # Process each text replacement
        for placeholder, replacement_text in text_replacements.items():
            replaced_count = 0
            
            # Search through table cells
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_num = row_idx * len(table.columns) + col_idx + 1
                    
                    # Check each paragraph in the cell
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if placeholder in paragraph.text:
                            print(f"\nProcessing Cell {cell_num} [row {row_idx}, col {col_idx}]:")
                            print(f"  Found '{placeholder}' in paragraph {para_idx}")
                            
                            # Handle multiline replacement
                            lines = replacement_text.split('\n')
                            
                            # Replace the placeholder in the first line
                            first_line = lines[0] if lines else ""
                            
                            # Replace text in runs to preserve formatting
                            replaced = False
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, first_line)
                                    replaced = True
                                    break
                            
                            if not replaced:
                                paragraph.text = paragraph.text.replace(placeholder, first_line)
                            
                            print(f"  Line 1: '{placeholder}' -> '{first_line}'")
                            
                            # Add additional lines as new paragraphs with same formatting
                            if len(lines) > 1:
                                for i, additional_line in enumerate(lines[1:], 2):
                                    if additional_line.strip():  # Only add non-empty lines
                                        new_para = cell.add_paragraph()
                                        
                                        # Try to copy paragraph style/formatting
                                        try:
                                            if hasattr(paragraph, 'style') and paragraph.style:
                                                new_para.style = paragraph.style
                                        except:
                                            pass
                                        
                                        # Add the text run
                                        new_run = new_para.add_run(additional_line)
                                        
                                        # Try to copy character formatting from original run
                                        try:
                                            if paragraph.runs:
                                                original_run = paragraph.runs[0]
                                                # Copy basic formatting properties
                                                if hasattr(original_run, 'font'):
                                                    new_run.font.name = original_run.font.name
                                                    new_run.font.size = original_run.font.size
                                                    new_run.font.bold = original_run.font.bold
                                                    new_run.font.italic = original_run.font.italic
                                                    new_run.font.color.rgb = original_run.font.color.rgb
                                        except:
                                            pass
                                        
                                        print(f"  Line {i}: Added '{additional_line}' (with formatting)")
                            
                            replaced_count += 1
            
            if replaced_count > 0:
                print(f"✅ Replaced '{placeholder}' in {replaced_count} location(s)")
            else:
                print(f"⚠️  Placeholder '{placeholder}' not found")
        
        # Save the document
        doc.save(output_path)
        print(f"\nDocument saved as: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return False

def replace_all_images_with_dummies(doc_path, output_path, dummy_dir="dummy_images"):
    """
    Replace ALL images in document with numbered dummy images in order.
    
    Args:
        doc_path (str): Path to input document
        output_path (str): Path to save modified document
        dummy_dir (str): Directory containing dummy images
    """
    try:
        if doc_path != output_path:
            shutil.copy2(doc_path, output_path)
        
        doc = Document(output_path)
        
        # Get list of dummy images
        dummy_files = sorted([f for f in os.listdir(dummy_dir) if f.endswith('.jpg')])
        if not dummy_files:
            print(f"No dummy images found in {dummy_dir}")
            return False
        
        print(f"Found {len(dummy_files)} dummy images")
        
        image_count = 0
        
        # Check paragraphs for images
        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                # Check for both modern and legacy image formats
                blips = run._element.xpath('.//a:blip')
                picts = run._element.xpath('.//w:pict')
                
                if blips or picts:
                    if image_count < len(dummy_files):
                        dummy_file = os.path.join(dummy_dir, dummy_files[image_count])
                        
                        # Get current dimensions if possible
                        current_width = None
                        try:
                            drawings = run._element.xpath('.//w:drawing')
                            for drawing in drawings:
                                inlines = drawing.xpath('.//wp:inline') + drawing.xpath('.//wp:anchor')
                                for inline in inlines:
                                    extents = inline.xpath('.//wp:extent')
                                    if extents:
                                        cx = int(extents[0].get('cx', 0))
                                        if cx > 0:
                                            current_width = cx / 914400
                                            break
                        except:
                            pass
                        
                        # Remove the existing run
                        paragraph._element.remove(run._element)
                        
                        # Add new dummy image
                        new_run = paragraph.add_run()
                        width = Inches(current_width) if current_width else Inches(2.0)
                        new_run.add_picture(dummy_file, width=width)
                        
                        image_count += 1
                        print(f"Image {image_count}: Replaced paragraph image with {dummy_files[image_count-1]}")
        
        # Check tables for images
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_num = row_idx * len(table.columns) + col_idx + 1
                    
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            # Check for both modern and legacy image formats
                            blips = run._element.xpath('.//a:blip')
                            picts = run._element.xpath('.//w:pict')
                            
                            if blips or picts:
                                if image_count < len(dummy_files):
                                    dummy_file = os.path.join(dummy_dir, dummy_files[image_count])
                                    
                                    # Get current dimensions if possible
                                    current_width = None
                                    try:
                                        drawings = run._element.xpath('.//w:drawing')
                                        for drawing in drawings:
                                            inlines = drawing.xpath('.//wp:inline') + drawing.xpath('.//wp:anchor')
                                            for inline in inlines:
                                                extents = inline.xpath('.//wp:extent')
                                                if extents:
                                                    cx = int(extents[0].get('cx', 0))
                                                    if cx > 0:
                                                        current_width = cx / 914400
                                                        break
                                    except:
                                        pass
                                    
                                    # Remove the existing run
                                    paragraph._element.remove(run._element)
                                    
                                    # Add new dummy image
                                    new_run = paragraph.add_run()
                                    width = Inches(current_width) if current_width else Inches(2.0)
                                    new_run.add_picture(dummy_file, width=width)
                                    
                                    image_count += 1
                                    print(f"Image {image_count}: Replaced cell {cell_num} image with {dummy_files[image_count-1]}")
        
        doc.save(output_path)
        print(f"\nSuccessfully replaced {image_count} images in {output_path}")
        return image_count > 0
        
    except Exception as e:
        print(f"Error replacing images: {str(e)}")
        return False

def replace_text_with_image(doc_path, output_path, text_to_replace, image_path, width_inches=2.0):
    """
    Replace specified text with an image in a Word document.
    
    Args:
        doc_path (str): Path to the input document
        output_path (str): Path to save the modified document
        text_to_replace (str): Text to replace with image
        image_path (str): Path to the replacement image
        width_inches (float): Width of the image in inches
    """
    try:
        if not os.path.exists(image_path):
            print(f"Error: Image file '{image_path}' not found.")
            return False
        
        doc = Document(output_path)  # Work on the already copied document
        
        replaced_count = 0
        
        # Search through all paragraphs
        for paragraph in doc.paragraphs:
            if text_to_replace in paragraph.text:
                # Clear the paragraph
                paragraph.clear()
                # Add the image
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                replaced_count += 1
                print(f"Replaced '{text_to_replace}' in paragraph with image: {image_path}")
        
        # Search through table cells
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_num = row_idx * len(table.columns) + col_idx + 1
                    for paragraph in cell.paragraphs:
                        if text_to_replace in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(width_inches))
                            replaced_count += 1
                            print(f"Replaced '{text_to_replace}' in cell {cell_num} with image: {image_path}")
        
        if replaced_count > 0:
            doc.save(output_path)
            return True
        else:
            print(f"Warning: Text '{text_to_replace}' not found in document")
            return False
        
    except Exception as e:
        print(f"Error replacing text with image: {str(e)}")
        return False

def main():
    """Main function to test both text and image replacement on Gafetes.docx"""
    input_file = "Gafetes.docx"
    output_file = "Gafetes_test_formatting.docx"
    
    print("=== Testing Gafetes.docx Processing (New Format) ===")
    print(f"Input: {input_file}")
    print(f"Output: {output_file}")
    
    # Step 1: Replace text placeholders with multiline data
    print("\n--- STEP 1: Text Placeholder Replacement ---")
    
    # Example data with different line counts per person
    text_data = {
        'text1': 'John Smith\nSoftware Engineer',  # 2 lines
        'text2': 'Maria Garcia\nProject Manager\n456 Oak Ave\nTown, ST 12345',  # 4 lines
        'text3': 'David Johnson\nData Analyst\n789 Pine Rd\nVillage, ST\nPhone: (555) 123-4567'  # 5 lines
    }
    
    print("Replacing text placeholders with:")
    for placeholder, text in text_data.items():
        lines = text.split('\n')
        print(f"  {placeholder}: {lines[0]} (+ {len(lines)-1} more lines)")
    
    text_success = replace_text_placeholders(input_file, output_file, text_data)
    
    if not text_success:
        print(f"\n❌ Text replacement failed!")
        return
    
    print(f"\n✅ Text replacement completed!")
    
    # Step 2: Replace image placeholders with dummy images
    print("\n--- STEP 2: Image Placeholder Replacement ---")
    print("Replacing 'image1', 'image2', 'image3' text with dummy images...")
    
    image_replacements = [
        ("image1", "dummy_images/dummy_01.jpg"),
        ("image2", "dummy_images/dummy_02.jpg"),
        ("image3", "dummy_images/dummy_03.jpg")
    ]
    
    image_success_count = 0
    for text_placeholder, dummy_image in image_replacements:
        success = replace_text_with_image(output_file, output_file, text_placeholder, dummy_image, width_inches=0.8)
        if success:
            image_success_count += 1
    
    if image_success_count > 0:
        print(f"\n✅ Image replacement completed! Replaced {image_success_count}/3 image placeholders")
    else:
        print(f"\n⚠️  No image placeholders found to replace")
    
    print(f"\n🎉 Processing complete! Check {output_file}")
    print("The document now has:")
    print("  - Replaced text placeholders with multiline dummy data")
    print(f"  - Replaced {image_success_count} image placeholders with dummy images")
    print("\nThis new approach is much more flexible!")
    print("- You can provide 2, 3, 4, or any number of lines per text placeholder")
    print("- Just use \\n to separate lines in your input strings")

if __name__ == "__main__":
    main()