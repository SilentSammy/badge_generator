from docx import Document
from docx.shared import Inches
import json
import shutil
import os
import pandas as pd
from docx2pdf import convert

def replace_text_in_document(doc, replacements):
    """
    Replace text in a Word document like Ctrl+H find-and-replace.
    Preserves formatting by working with runs instead of paragraph.text
    
    Args:
        doc: Document object
        replacements (dict): Dictionary with find:replace pairs
    """
    for find_text, replace_text in replacements.items():
        replaced_count = 0
        
        # Handle multiline replacement text
        if '\n' in replace_text:
            # For multiline text, use Word's line break
            replace_text = replace_text.replace('\n', '\r')
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                replaced_count += replace_in_paragraph(paragraph, find_text, replace_text)
        
        # Replace in all table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            replaced_count += replace_in_paragraph(paragraph, find_text, replace_text)
        
        if replaced_count > 0:
            print(f"✅ Replaced '{find_text}' in {replaced_count} location(s)")
        else:
            print(f"⚠️  Warning: '{find_text}' not found in document")

def replace_in_paragraph(paragraph, find_text, replace_text):
    """
    Replace text in a paragraph while preserving formatting by working with runs.
    Handles text that spans multiple runs by reconstructing the paragraph.
    
    Args:
        paragraph: Word paragraph object
        find_text (str): Text to find
        replace_text (str): Text to replace with
        
    Returns:
        int: Number of replacements made (0 or 1)
    """
    if find_text not in paragraph.text:
        return 0
    
    # Try to replace within individual runs first (preserves formatting perfectly)
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return 1
    
    # If text spans multiple runs, we need to reconstruct while preserving formatting
    if find_text in paragraph.text:
        # Store all run information (text + formatting)
        run_data = []
        for run in paragraph.runs:
            run_data.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font.name else None,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color.rgb else None
            })
        
        # Get the full text and perform replacement
        full_text = paragraph.text
        new_text = full_text.replace(find_text, replace_text)
        
        # Clear the paragraph
        paragraph.clear()
        
        # If the replacement is simple and fits in the first run's style, use that
        if len(run_data) > 0:
            new_run = paragraph.add_run(new_text)
            # Apply the formatting from the first run
            first_run = run_data[0]
            if first_run['bold'] is not None:
                new_run.bold = first_run['bold']
            if first_run['italic'] is not None:
                new_run.italic = first_run['italic']
            if first_run['underline'] is not None:
                new_run.underline = first_run['underline']
            if first_run['font_name']:
                new_run.font.name = first_run['font_name']
            if first_run['font_size']:
                new_run.font.size = first_run['font_size']
            if first_run['color']:
                new_run.font.color.rgb = first_run['color']
        else:
            # Fallback: just add the text
            paragraph.add_run(new_text)
        
        return 1
    
    return 0

def replace_images_in_document(doc, image_replacements, width_inches=0.8):
    """
    Replace text placeholders with images in a Word document.
    
    Args:
        doc: Document object
        image_replacements (dict): Dictionary with placeholder:image_path pairs
    """
    for placeholder, image_path in image_replacements.items():
        replaced_count = 0
        
        # Check if image file exists
        if not os.path.exists(image_path):
            print(f"⚠️  Warning: Image file '{image_path}' not found")
            continue
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                replaced_count += replace_text_with_image_in_paragraph(paragraph, placeholder, image_path, width_inches=width_inches)
        
        # Replace in all table cells  
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            replaced_count += replace_text_with_image_in_paragraph(paragraph, placeholder, image_path, width_inches=width_inches)
        
        if replaced_count > 0:
            print(f"✅ Replaced '{placeholder}' with image in {replaced_count} location(s)")
        else:
            print(f"⚠️  Warning: '{placeholder}' not found in document")

def replace_text_with_image_in_paragraph(paragraph, placeholder, image_path, width_inches=0.8):
    """
    Replace text placeholder with an image in a paragraph.
    
    Args:
        paragraph: Word paragraph object
        placeholder (str): Text to replace
        image_path (str): Path to replacement image
        width_inches (float): Width of the image in inches
        
    Returns:
        int: Number of replacements made (0 or 1)
    """
    if placeholder not in paragraph.text:
        return 0
    
    # Clear the paragraph and add the image
    paragraph.clear()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    return 1

def table_to_dicts(table, count_per_doc=30, image_column="Image", key_format="{}{:02d}"):
    """
    Split a DataFrame into groups and generate text/image replacement dictionaries.
    
    Args:
        table: pandas DataFrame with attendee data
        count_per_doc (int): Number of attendees per document
        image_column (str): Column name containing image paths
        key_format (str): Format string for generating keys (e.g., "name{:02d}")
        
    Returns:
        list: List of tuples (text_dict, image_dict) for each document group
    """
    result_groups = []
    
    # Split DataFrame into chunks of count_per_doc rows
    for i in range(0, len(table), count_per_doc):
        chunk = table.iloc[i:i + count_per_doc]
        
        text_dict = {}
        image_dict = {}
        
        # Generate dictionaries for this chunk
        for idx, (_, row) in enumerate(chunk.iterrows(), 1):
            # Generate keys for text columns (excluding image column)
            for column in table.columns:
                if column != image_column and column != 'id':  # Skip image and id columns
                    key = key_format.format(column, idx)
                    text_dict[key] = str(row[column])
            
            # Generate key for image column
            if image_column in table.columns:
                image_key = key_format.format(image_column, idx)
                image_dict[image_key] = str(row[image_column])
        
        result_groups.append((text_dict, image_dict))
    
    return result_groups

def main():
    """Load CSV data and generate multiple Word documents using profiles"""
    profiles = [
        {
            "name": "Standard Badges",
            "csv_file": "Atendees.csv",
            "template_file": "Gafetes.docx",
            "count_per_doc": 30,
            "output_prefix": "Gafetes/Gafetes",
            "image_column": "Image",
            "key_format": "{}{:02d}",
            "image_width_inches": 0.8
        },
        # Add more profiles here as needed
        {
            "name": "Entry Passes",
            "csv_file": "Pases.csv",
            "template_file": "Pases.docx",
            "count_per_doc": 12,
            "output_prefix": "Pases/Pases",
            "image_column": "Image",
            "key_format": "{}{:02d}",
            "image_width_inches": 1.75
        },
    ]
    
    # Select which profile to use (change index to switch profiles)
    profile = profiles[1]
    
    # Extract profile settings
    csv_file = profile["csv_file"]
    template_file = profile["template_file"]
    count_per_doc = profile["count_per_doc"]
    output_prefix = profile.get("output_prefix", "Document")
    profile_name = profile.get("name", "Unnamed Profile")
    image_column = profile.get("image_column", "Image")
    key_format = profile.get("key_format", "{}{:02d}")
    
    print(f"=== {profile_name} ===")
    print(f"Reading: {csv_file}")
    print(f"Template: {template_file}")
    print(f"Attendees per document: {count_per_doc}")
    print(f"Output prefix: {output_prefix}")
    print(f"Image column: {image_column}")
    print(f"Key format: {key_format}")
    
    # Load CSV data
    try:
        df = pd.read_csv(csv_file)
        print(f"✅ Loaded {len(df)} attendees from CSV")
        print(f"Columns: {list(df.columns)}")
        print()
    except Exception as e:
        print(f"❌ Error loading CSV: {e}")
        return
    
    # Check if template exists
    if not os.path.exists(template_file):
        print(f"❌ Template file '{template_file}' not found")
        return
    
    # Split into document groups
    groups = table_to_dicts(df, count_per_doc=count_per_doc, image_column=image_column, key_format=key_format)
    print(f"📄 Generating {len(groups)} documents...")
    print()
    
    # Generate documents for each group
    for i, (text_dict, image_dict) in enumerate(groups):
        output_file = f"{output_prefix}_{i:02d}.docx"
        
        print(f"🔄 Processing document {i+1}/{len(groups)}: {output_file}")
        
        try:
            # Copy template to output file
            shutil.copy2(template_file, output_file)
            doc = Document(output_file)
            
            # Perform text replacements
            print(f"   📝 Performing {len(text_dict)} text replacements...")
            replace_text_in_document(doc, text_dict)
            
            # Perform image replacements (fix paths by removing leading slash)
            fixed_image_dict = {}
            for key, path in image_dict.items():
                # Remove leading slash if present
                fixed_path = path.lstrip('/')
                fixed_image_dict[key] = fixed_path
            
            print(f"   🖼️  Performing {len(fixed_image_dict)} image replacements...")
            replace_images_in_document(doc, fixed_image_dict, width_inches=profile.get("image_width_inches", 0.8))
            
            # Save document
            doc.save(output_file)
            print(f"   ✅ Saved: {output_file}")
            
            # Convert to PDF
            pdf_file = f"{output_prefix}_{i:02d}.pdf"
            print(f"   📄 Converting to PDF: {pdf_file}")
            try:
                convert(output_file, pdf_file)
                print(f"   ✅ PDF saved: {pdf_file}")
            except Exception as pdf_error:
                print(f"   ⚠️  PDF conversion failed: {pdf_error}")
            
        except Exception as e:
            print(f"   ❌ Error processing {output_file}: {e}")
        
        print()
    
    print(f"🎉 Profile '{profile_name}' complete! Generated {len(groups)} documents with PDF exports.")

if __name__ == "__main__":
    main()